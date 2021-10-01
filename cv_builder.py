from docx import Document
from docx.shared import Inches
import pyttsx3



def speak(text): 
    pyttsx3.speak(text) 


document = Document()


# profile picture
document.add_picture('teimoor.jpg',
                     width=Inches(2.0))

# name, phone number and email details
name = input("What is your name? ")
speak('Hello ' + name + ' how are you today?')

speak("What is your phone number? ")
phone_number = input("What is your phone number? ")
email = input("What is your email ? ")

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# about me
document.add_heading('About me')
about_me = input("Tell me about yourself? ")
document.add_paragraph(about_me)


# Work experiences
document.add_heading('Work Experiences')
p = document.add_paragraph()

company = input("Enter company: ")
from_date = input('From Date: ')
to_date = input("To Date: ")

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input("Describe your experience at " + company + " ")

p.add_run(experience_details)


# More experiences
while True:
    has_more_experiences = input("Do you have more experiences? (Yes or No) ")

    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input("Enter company: ")
        from_date = input('From Date: ')
        to_date = input("To Date: ")

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input("Describe your experience at " + company + " ")

        p.add_run(experience_details)
    else:
        break


# skills section 
document.add_heading('Skills') 
skill = input("Enter skills") 
p = document.add_paragraph(skill, style="List Bullet") 

while True: 
    has_more_skills = input("Do you have more skills? (y/n) ") 
    if has_more_skills.lower() == 'y': 
        skill = input("Enter skill: ") 
        p = document.add_paragraph(skill, style="List Bullet") 
    else:
        break


# footer 
section  = document.sections[0] 
footer = section.footer 
p = footer.paragraphs[0] 
p.text = "Cv generated using python programming and with the python-docx module" 

document.save('cv.docx')

